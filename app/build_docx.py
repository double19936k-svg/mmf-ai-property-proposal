from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Microsoft YaHei"
NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(46, 117, 182)
INK = RGBColor(32, 42, 56)
MUTED = RGBColor(94, 107, 120)
LIGHT = "F2F4F7"
WHITE = "FFFFFF"


def set_run(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def configure_style(style, size, color, bold, before, after, line_spacing, align=WD_ALIGN_PARAGRAPH.LEFT):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run(run, size=9, color=MUTED)


def add_title_block(doc, artifact, brief):
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(8)
    set_run(kicker.add_run(brief.get("project_name", "物业项目")), size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    set_run(title.add_run(artifact.get("title", brief.get("scenario", "物业方案章节"))), size=25, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run(subtitle.add_run(artifact.get("subtitle", f"{brief.get('project_name', '')}｜{brief.get('scenario', '')}")), size=12, color=MUTED)


def add_table(doc, table_data):
    columns = table_data.get("columns") or []
    rows = table_data.get("rows") or []
    if not columns or not rows or any(len(row) != len(columns) for row in rows):
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for _ in rows:
        table.add_row()
    total = 9360
    base = total // len(columns)
    widths = [base] * len(columns)
    widths[-1] += total - sum(widths)
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    values = [columns, *rows]
    for r_index, row_values in enumerate(values):
        for c_index, value in enumerate(row_values):
            cell = table.rows[r_index].cells[c_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(cell, "17365D" if r_index == 0 else (LIGHT if r_index % 2 else WHITE))
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=9.5, bold=r_index == 0, color=RGBColor(255, 255, 255) if r_index == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_docx.py input.json output.docx")
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    payload = json.loads(source.read_text(encoding="utf-8-sig"))
    brief = payload["brief"]
    artifact = payload["artifact"]
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], 11, INK, False, 0, 8, 1.333, WD_ALIGN_PARAGRAPH.JUSTIFY)
    configure_style(doc.styles["Heading 1"], 16, BLUE, True, 18, 10, 1.0)
    configure_style(doc.styles["Heading 2"], 13, BLUE, True, 12, 6, 1.0)
    configure_style(doc.styles["Heading 3"], 12, NAVY, True, 8, 4, 1.0)
    configure_style(doc.styles["List Bullet"], 11, INK, False, 0, 4, 1.208)

    header_p = section.header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    set_run(header_p.add_run(f"{brief.get('project_name', '')}｜{brief.get('scenario', '')}"), size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    add_title_block(doc, artifact, brief)
    for lead in artifact.get("lead", []):
        p = doc.add_paragraph(str(lead))
        p.paragraph_format.space_after = Pt(10)
        for run in p.runs:
            set_run(run, size=11, color=INK)

    for section_data in artifact.get("sections", []):
        heading = str(section_data.get("heading", "")).strip()
        if not heading:
            continue
        doc.add_heading(heading, level=1)
        for text in section_data.get("paragraphs", []):
            p = doc.add_paragraph(str(text))
            for run in p.runs:
                set_run(run, size=11, color=INK)
        for text in section_data.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            set_run(p.add_run(str(text)), size=11, color=INK)
        if isinstance(section_data.get("table"), dict):
            add_table(doc, section_data["table"])

    doc.core_properties.title = artifact.get("title", "物业方案章节初稿")
    doc.core_properties.subject = brief.get("scenario", "物业方案")
    doc.core_properties.author = "物业服务方案编制团队"
    doc.save(output)


if __name__ == "__main__":
    main()
